import re
import itertools
from fuzzywuzzy import fuzz
from docx import Document
import enchant
import time
from collections import Counter
import nltk
from nltk.collocations import *
from nltk.corpus import stopwords
from nltk import pos_tag, word_tokenize
import argparse
import os
from argparse import ArgumentParser
import pathlib
import shutil


start_time = time.time()
d = enchant.Dict("en")
stopWords = set(stopwords.words('english'))

punctuationset = "!#$%&'()*+, “-./‘—:;<=>’?@[\]^\"_`{|}~”„…"
trigram_measures = nltk.collocations.TrigramAssocMeasures()
fourgram_measures = nltk.collocations.QuadgramAssocMeasures()

class Manu:
    def __init__(self, filename):
        self.filename = filename
        self.AllWordsNew = []
        self.WordCount = 0
        self.CarvedWordsNew = []
        self.uniquecarvedwords = set()
        self.uniqueallwords = set()
        self.propwords = []
        self.freqcount = Counter()
        self.addvalues()

    def addvalues(self):
        for paragraph in Document(self.filename).paragraphs:
            self.AllWordsNew.append(paragraph.text.split())
        self.AllWordsNew = list(itertools.chain(*self.AllWordsNew))
        self.WordCount = len(self.AllWordsNew)
        text_file = open("mtempman.txt", "w", encoding="UTF-8")
        for x in range(len(self.AllWordsNew)):
            text_file.write(self.AllWordsNew[x] + " ")
            self.CarvedWordsNew.append(self.word_carver(self.AllWordsNew[x]))
        text_file.close()

        for word in self.AllWordsNew:
            self.uniqueallwords.add(word)

        for word in self.CarvedWordsNew:
            self.uniquecarvedwords.add(word)

        self.uniquecarvedwords.remove("")

        for word in self.uniquecarvedwords:
            if not d.check(word):
                self.propwords.append(word)

        self.propwords = sorted(self.propwords, key=str.casefold)
        self.freqcount = Counter(self.AllWordsNew)

    def word_carver(self,uncarved_word):
        if any(map(str.isdigit, uncarved_word)):
            return ""
        elif re.search(r'([A-Z,a-z])+(\.)+([A-Z,a-z])+(\.)',uncarved_word) != None:
            return re.search(r'([A-Z,a-z])+(\.)+([A-Z,a-z])+(\.)',uncarved_word).group()
        elif uncarved_word in punctuationset:
            return ""
        elif uncarved_word[0] in punctuationset:
            uncarved_word = uncarved_word[1:]
            return self.word_carver(uncarved_word)
        elif uncarved_word[-1] in punctuationset:
            uncarved_word = uncarved_word[:-1]
            return self.word_carver(uncarved_word)
        else:
            return uncarved_word

    def fuzzymatch(self):
        wordcomps = []
        finalproduct = []
        for word, OtherWord in itertools.combinations(self.propwords, 2):
            wordRatio = fuzz.token_set_ratio(word, OtherWord)
            if word.isupper() == False and OtherWord.isupper() == False:
                if 99 > wordRatio > 75:
                    wordcomps.append([word, OtherWord, wordRatio])

        for word in wordcomps:
            finalproduct.append(word[0] + " & " + word[1])

        return(finalproduct)


    def trigramcounter(self):
        finalproduct=[]
        with open("mtempman.txt", "r", encoding="UTF-8") as text_file:
            tokens = nltk.wordpunct_tokenize(text_file.read())
            finder = TrigramCollocationFinder.from_words(tokens)
            finder.apply_word_filter(lambda w: w in punctuationset)
            finder.apply_freq_filter(3)
            len(finder.score_ngrams(trigram_measures.raw_freq))
            scored = finder.score_ngrams(trigram_measures.raw_freq)
            sorted(gram for gram, score in scored)
            text_file.close()

        for trigram, score in scored:
            line = (trigram[0] + " " + trigram[1] + " " + trigram[2])
            line = str(line)
            finalproduct.append(line)
        if len(finalproduct) > 30:
            return finalproduct[:30]
        else:
            return finalproduct


    def quadgramcounter(self):
        finalproduct = []
        with open("mtempman.txt", "r", encoding="UTF-8") as text_file:
            tokens = nltk.wordpunct_tokenize(text_file.read())
            finder = QuadgramCollocationFinder.from_words(tokens)
            finder.apply_word_filter(lambda w: w in punctuationset)
            finder.apply_freq_filter(3)
            len(finder.score_ngrams(fourgram_measures.raw_freq))
            scored = finder.score_ngrams(fourgram_measures.raw_freq)
            sorted(gram for gram, score in scored)
            text_file.close()

        for trigram, score in scored:
            line = (trigram[0] + " " + trigram[1] + " " + trigram[2] + " " + trigram[3])
            line = str(line)
            finalproduct.append(line)
        if len(finalproduct) > 30:
            return finalproduct[:30]
        else:
            return finalproduct

    def possessivefinder(self):
        possessives = set()
        finalproduct = []
        for word in self.AllWordsNew:
            if "’s" in word:
                possessives.add(word)
            elif "s’" in word:
                possessives.add(word)

        for pos in possessives:
            finalproduct.append(pos)
        return (finalproduct)

    def splitinfs(self):
        splits = []
        POS = []
        for word, token in (pos_tag(self.AllWordsNew)):
            POS.append([word, token])

        for x in range(len(POS) - 4):
            if POS[x][1] == "TO":
                if POS[x + 1][1] == "RB":
                    if POS[x + 2][1] == "VB":
                        splits.append(POS[x][0] + " " + POS[x + 1][0] + " " + POS[x + 2][0])
                elif POS[x + 1][1] == "RB":
                    if POS[x + 3][1] == "VB":
                        splits.append(POS[x][0] + " " + POS[x + 1][0] + " " + POS[x + 2][0] + " " + POS[x + 3][0])
        return (splits)

    def freqcounter(self):
        finalproduct=[]
        for word, counter in self.freqcount.most_common():
            if counter > (self.WordCount * .001):
                if word not in stopWords and word.lower() not in stopWords and word != " ":
                    finalproduct.append(word + " " + str(counter))
        if len(finalproduct) > 30:
            return finalproduct[:30]
        else:
            return finalproduct

    def ampersandtest(self):
        ampgroup = set()
        andgroup = set()
        for paragraph in Document(self.filename).paragraphs:
            Ampersand = re.search(r'([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+)?([-])?([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+ ){0,2}([&] )([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+) (et al. )?[1-2]+(\d{3})+[a-zA-Z]?', paragraph.text)
            Spelled = re.search(r'([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+)?([-])?([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+ ){0,2}(\band\b )([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+) [1-2]+(\d{3})+[a-zA-Z]?',
                paragraph.text)
            if Ampersand is not None:
                ampgroup.add(Ampersand.group())
            if Spelled is not None:
                andgroup.add(Spelled.group())
        ampgroup = list(ampgroup)
        andgroup = list(andgroup)
        if len(ampgroup) and len(andgroup) > 0:
            if len(ampgroup) == len(andgroup):
                return(["Citations with two authors use ampersand/and interchangably"])
            elif len(ampgroup) < len(andgroup):
                return(ampgroup)
            else:
                return(andgroup)
        else:
            return ([])

class Biblo(Manu):
    def __init__(self, filename, style):
        self.filename = filename
        self.style = style
        self.AllWordsNew = []
        self.WordCount = 0
        self.CarvedWordsNew = []
        self.uniquecarvedwords = set()
        self.uniqueallwords = set()
        self.propwords = []
        self.freqcount = Counter()
        self.addvalues()
        self.extractedauthors = []
        self.addbibvalues()


    def addbibvalues(self):
        self.extractedauthors = self.authorextractTest()

    def dashfixer(self, ListofAuthors):
        finalauthors = []
        for x in range(len(ListofAuthors)):
            if re.search(r'[\u002D\u058A\u05BE\u1400\u1806\u2010-\u2015\u2E17\u2E1A\u2E3A\u2E3B\u2E40\u301C\u3030\u30A0\uFE31\uFE32\uFE58\uFE63\uFF0D]', ListofAuthors[x][0][0]) is not None:
                ListofAuthors[x][0][0] = ListofAuthors[x - 1][0][0]
            finalauthors.append([ListofAuthors[x][0][0], ListofAuthors[x][1][0]])
        return (finalauthors)


    def authorpatternsearcher(self,patternlist, text):
        for pattern in patternlist:
            if pattern is not None and pattern.group() in text:
                return True

    def authorpatternwriter(self, authorcount, authornames, yearposition):
        group1 = ["AJA","HESPERIA"]
        group2 =["ANTIQUITY"]
        group3 = ["JOURNAL OF MATERIAL CULTURE"]

        if self.style in group3:
            authorname = re.search(r'^([^(]*)', authornames)
        else:
            authorname = re.search(r'^([^,]*)', authornames)

        if authorcount == 1:
            return [[authorname.group()], [yearposition]]
        elif authorcount == 2:
            if self.style in group2:
                authorname2 = re.search(r'((?<=& ).*)|((?<=\band\b ).*).*', authornames)
                secondauthorcleaned = re.search(r'^([^,])+', authorname2.group())
                return [[authorname.group() + " & " + secondauthorcleaned.group()], [yearposition]]
            elif self.style in group1:
                authorname2 = re.search(r'((?<=& ).*)|((?<=\band\b ).*).*', authornames)
                secondauthorcleaned = re.search(r'(?<=. )(.*)(?=\.)', authorname2.group())
                return [[authorname.group() + " and " + secondauthorcleaned.group()], [yearposition]]
            elif self.style in group3:
                authorname2 = re.search(r',.*$', authornames)
                secondauthorcleaned = authorname2
                return [[authorname.group() + " and " + secondauthorcleaned.group()], [yearposition]]
        elif authorcount > 2:
            return [[authorname.group() + " et al."], [yearposition]]
        else:
            return [["Possible Issue with "+authornames],["Possible Issue with"]]


    def has_numbers(self,inputString):
        return any(char.isdigit() for char in inputString)

    def authorcounter(self,authornames):
        authorcount = authornames.count(",")
        if ", ed." in authornames:
            authorcount = authorcount - 1
        elif ", Jr." in authornames:
            authorcount = authorcount - 1
        elif "et al" in authornames:
            authorcount = 1
        return(authorcount)

    def yearextractor(self,text):
        if "Forthcoming" in text:
            yearposition = "Forthcoming"
        else:
            yearposition = re.search(r'\d{4}(\w)?', text)
            if yearposition is not None:
                yearposition = yearposition.group()
            else:
                yearposition = "N.D."
        return(yearposition)

    def authorextractTest(self):
        returnedauthors = []
        authorerrors=[]
        group1 = ["AJA","HESPERIA","ANTIQUITY"]
        group2 =["JOURNAL OF MATERIAL CULTURE"]
        for paragraph in Document(self.filename).paragraphs:
            if len(paragraph.text) < 5:
                continue
            else:
                yearposition = self.yearextractor(paragraph.text)

                if re.search(r'[\u002D\u058A\u05BE\u1400\u1806\u2010-\u2015\u2E17\u2E1A\u2E3A\u2E3B\u2E40\u301C\u3030\u30A0\uFE31\uFE32\uFE58\uFE63\uFF0D]',
                        paragraph.text[0]) is not None:
                    actualauthor = ["—"]
                    returnedauthors.append([actualauthor, [yearposition]])
                elif self.style in group1:
                    actualauthor = re.search(rf'.+?(?={yearposition})', paragraph.text)
                    if actualauthor is not None:
                        authornames = actualauthor.group()
                        authornames = authornames.strip()
                        authorcount = self.authorcounter(authornames)
                        returnedauthors.append(self.authorpatternwriter(authorcount, authornames, yearposition))
                    else:
                        authorerrors.append(paragraph.text+" Missing Date?")
                elif self.style in group2:
                    actualauthor = re.search(r'^([^(]*)', paragraph.text)
                    authornames = actualauthor.group()
                    authornames = authornames.strip()
                    authorcount = self.authorcounter(authornames)
                    returnedauthors.append(self.authorpatternwriter(authorcount, authornames, yearposition))
        returnedauthors = self.dashfixer(returnedauthors)

        return (returnedauthors)

class texdoc(Manu):
    def __init__(self, filename):
        self.filename = filename
        self.AllWordsNew = []
        self.CarvedWords = []
        self.WordCount = 0
        self.uniquecarvedwords = set()
        self.uniqueallwords = set()
        self.propwords = []
        self.freqcount = Counter()
        self.addvalues()


    def addvalues(self):
        target =os.getcwd()+"\mtempman.txt"
        shutil.copyfile(self.filename, target)
        with open(self.filename, encoding="utf8") as f:
            for line in f:
                if line[0] != "%":
                    latcommand = re.search(r'\$?\\(.*?)}', line)
                    if latcommand is not None:
                        line2 = re.sub(r'\$?\\(.*?)}',"",line)
                        self.AllWordsNew.append(line2.split())
                    else:
                        self.AllWordsNew.append(line.split())

        self.AllWordsNew = list(itertools.chain(*self.AllWordsNew))

        for word in self.AllWordsNew:
            self.CarvedWords.append(self.word_carver(word))

        self.WordCount = len(self.CarvedWords)


        for word in self.AllWordsNew:
            self.uniqueallwords.add(word)

        for word in self.CarvedWords:
            self.uniquecarvedwords.add(word)

        self.uniquecarvedwords.remove("")

        for word in self.uniquecarvedwords:
            if not d.check(word):
                self.propwords.append(word)

        self.propwords = sorted(self.propwords, key=str.casefold)
        self.freqcount = Counter(self.AllWordsNew)

class BasePublisher:
    def __init__(self, Manu):
        self.Manu = Manu
        self.finaloutput()

    def functionprinter(self,func,title,f):
        if len(func) > 0:
            f.write('\n'+ title + '\n')
            for x in func:
                f.write(x + '\n')

    def finaloutput(self):
        with open('Old Files/output.txt', 'w', encoding='utf-8') as f:
            f.write("Manuscript Wordcount: " + str(self.Manu.WordCount) + '\n')
            self.functionprinter(self.Manu.fuzzymatch(),"Unique Words With High Similarities in Manuscript",f)
            self.functionprinter(self.Manu.possessivefinder(), "Possessives in Manuscript", f)
            self.functionprinter(self.Manu.splitinfs(), "Split Infinitives", f)
            self.functionprinter(self.Manu.propwords, "All Proper Nouns/Non-Dict in Manuscript", f)
            self.functionprinter(self.Manu.freqcounter(), "All Frequencies", f)
            self.functionprinter(self.Manu.quadgramcounter(), "Common Four Word Phrases", f)
            self.functionprinter(self.Manu.trigramcounter(), "Common Three Word Phrases", f)
        f.close()

class ShortDocPublisher(BasePublisher):
    def __init__(self, Manu):
        self.Manu = Manu
        self.finaloutput()

    def functionprinter(self,func,title,f):
        if len(func) > 0:
            f.write('\n'+ title + '\n')
            for x in func:
                f.write(x + '\n')




class Publisher(BasePublisher):
    def __init__(self, Manu, Biblo):
        self.Manu = Manu
        self.Biblo = Biblo
        self.finaloutput()

    def crossdocfuzzymatch(self):
        wordcomps = []
        for word in self.Manu.propwords:
            for OtherWord in self.Biblo.propwords:
                if 99 > fuzz.token_set_ratio(word, OtherWord) > 75:
                    wordcomps.append(word + " & " + OtherWord)

        return wordcomps

    def authorpattern(self):
        authormistakes = []
        finalproduct = []
        for author in self.Biblo.extractedauthors:
            authorstatus = False
            for par in Document(self.Manu.filename).paragraphs:  # to extract the whole text
                SearchAuthor = author[0]
                AltSearchAuthor = SearchAuthor + "’s"
                SearchYear = author[1]
                magicpattern = re.search(SearchAuthor + r'([^a-zA-Z]+\w?){0,5}' + SearchYear, par.text)
                magicpattern2 = re.search(AltSearchAuthor + r'([^a-zA-Z]+\w?){0,5}' + SearchYear, par.text)

                if self.Biblo.authorpatternsearcher([magicpattern, magicpattern2], par.text) is True:
                    authorstatus = True

            if authorstatus == False:
                authormistakes.append(author)

        for mistake in authormistakes:
            finalproduct.append(mistake[0] + " " + mistake[1])
        return (finalproduct)

    def forgottencitation(self):
        possibleauthors = set()

        for paragraph in Document(self.Manu.filename).paragraphs:
            PossibleYear = re.search(r'([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+)?([-])?([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+ ){0,2}([&] )?(\band\b )?([A-Z][A-Za-zÀ-ÖØ-öø-ÿ]+) (et al. )?[1-2]+(\d{3})+[a-zA-Z]?', paragraph.text)
            if PossibleYear is not None and PossibleYear.group() in paragraph.text:
                if PossibleYear.group()[0].isupper():
                    if PossibleYear.group().isupper():
                        pass
                    elif PossibleYear.group().count(" ") > 1 and "See" in PossibleYear.group():
                        pass
                    elif "In " in PossibleYear.group():
                        pass
                    else:
                        possibleauthors.add(PossibleYear.group())

        for x in self.Biblo.extractedauthors:
            confirmedauthor = x[0] + " " + x[1]
            for author in possibleauthors:
                if author in confirmedauthor:
                    possibleauthors.remove(author)
                    break
        return possibleauthors

    def fuzzymatch(self):
        wordcomps = []
        for word, OtherWord in itertools.combinations(self.propwords, 2):
            wordRatio = fuzz.token_set_ratio(word, OtherWord)
            if 99 > wordRatio > 75:
                wordcomps.append([word, OtherWord, wordRatio])
        return wordcomps

    def functionprinter(self,func,title,f):
        if len(func) > 0:
            f.write('\n'+ title + '\n')
            for x in func:
                f.write(x + '\n')

    def errorfunctionprinter(self,func,title,errortitle,f):
        if len(func) > 0:
            f.write('\n'+ title + '\n')
            for x in func[0]:
                f.write(x + '\n')
            f.write('\n' + errortitle + '\n')
            for x in func[1]:
                f.write(x + '\n')

    def finaloutput(self):

        with open('Old Files/output.txt', 'w', encoding='utf-8') as f:
            f.write("Manuscript Wordcount: " + str(self.Manu.WordCount) + '\n')
            self.functionprinter(self.Manu.ampersandtest(), "Ampersand Inconsistency", f)
            self.functionprinter(self.Manu.fuzzymatch(),"Unique Words With High Similarities in Manuscript",f)
            self.functionprinter(self.Biblo.fuzzymatch(), "Unique Words With High Similarities in Bibliography", f)
            self.functionprinter(self.authorpattern(),"Appear to be in Bibliography, but not in Manuscript",f)
            self.functionprinter(self.forgottencitation(), "Appear to be citations, but not in Bibliography", f)
            self.functionprinter(self.crossdocfuzzymatch(), "Unique Words With High Similarities in Manuscript and Bibliography", f)
            self.functionprinter(self.Manu.possessivefinder(), "Possessives in Manuscript", f)
            self.functionprinter(self.Manu.splitinfs(), "Split Infinitives", f)
            self.functionprinter(self.Manu.propwords, "All Proper Nouns/Non-Dict in Manuscript", f)
            self.functionprinter(self.Biblo.propwords, "All Proper Nouns/Non-Dict in Bibliography", f)
            self.functionprinter(self.Manu.freqcounter(), "All Frequencies", f)
            self.functionprinter(self.Manu.quadgramcounter(), "Common Four Word Phrases", f)
            self.functionprinter(self.Manu.trigramcounter(), "Common Three Word Phrases", f)
        f.close()



def validate_file(f):
    if not os.path.exists(f):
        raise argparse.ArgumentTypeError("{0} does not exist".format(f))
    return f


if __name__ == '__main__':
    parser = ArgumentParser(description="Read file form Command line.")
    parser.add_argument("-m", "--manuscript", dest="manfilename", required=True, type=validate_file,
                            help="input file", metavar="manuscript path")
    parser.add_argument("-b", "--bibliography", dest="bibfilename", required=False, type=validate_file,
                            help="input file", metavar="bibliography path")
    parser.add_argument("-bt", "--bibtype", dest="bibtype", required=False, type=str,
                            help="input file", metavar="Bibliography Type")
    args = parser.parse_args()
    manuscript_extension = str(pathlib.Path(args.manfilename).suffix)
    if args.bibfilename is not None:
        Publisher(Manu(args.manfilename),Biblo(args.bibfilename, args.bibtype))
    elif manuscript_extension == ".docx":
        ShortDocPublisher(Manu(args.manfilename))
    elif manuscript_extension == ".tex":
        ShortDocPublisher(texdoc(args.manfilename))

    os.remove("mtempman.txt")











